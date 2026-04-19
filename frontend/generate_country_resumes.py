#!/usr/bin/env python3
"""
AI-Powered Multi-Country Resume Generator
==========================================
Features:
  - Parse an existing resume (DOCX / PDF / TXT) as source of truth
  - Google Gemini AI extracts structured data and generates country-tailored summaries
  - Beautifully formatted DOCX resumes for 8 countries:
        Canada | USA | UK | Europe (Europass) | Japan | Australia | UAE | Singapore
  - Optional job-description tailoring for keyword-matched output
  - Gracefully falls back to built-in data if AI key is not set
Requirements:
  pip install python-docx google-generativeai PyMuPDF
Environment variables:
  GEMINI_API_KEY       Google Gemini API key (required for AI features)
  RESUME_OUTPUT_DIR    Override default output folder
Usage:
  python generate_country_resumes.py                                    # built-in data, all countries
  python generate_country_resumes.py --source resume.docx               # parse your own resume
  python generate_country_resumes.py --source resume.pdf --jd jd.txt   # JD-tailored output
  python generate_country_resumes.py --countries canada usa uk           # specific countries
  python generate_country_resumes.py --enhance-bullets                   # AI-improved bullet points
  python generate_country_resumes.py --ai-summaries                      # country-tailored summaries
  python generate_country_resumes.py --list-countries                    # show supported countries
"""
import os
import sys
import json
import argparse
from copy import deepcopy
from dataclasses import dataclass, field
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
# == Config ===================================================================
OUTPUT_DIR     = os.environ.get("RESUME_OUTPUT_DIR", r"C:\Users\baljinders\Downloads\Resume")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
GEMINI_MODEL   = "gemini-2.0-flash"
SUPPORTED_COUNTRIES = ["canada", "usa", "uk", "europe", "japan", "australia", "uae", "singapore"]
# == Colour palette ===========================================================
NAVY  = RGBColor(0x1A, 0x3C, 0x6E)
LINK  = RGBColor(0x05, 0x63, 0xC1)
GRAY  = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x00, 0x00, 0x00)
# =============================================================================
#  DATA MODELS
# =============================================================================
@dataclass
class WorkEntry:
    company:      str
    location:     str
    dates:        str
    title:        str
    former_title: Optional[str] = None
    bullets:      list          = field(default_factory=list)
    industry:     Optional[str] = None
    company_size: Optional[str] = None
@dataclass
class EduEntry:
    institution: str
    degree:      str
    dates:       str
    location:    Optional[str] = None
@dataclass
class PatentEntry:
    title:       str
    year:        str
    description: str
@dataclass
class ResumeData:
    name:        str = ""
    email:       str = ""
    phone:       str = ""
    linkedin:    str = ""
    github:      str = ""
    website:     str = ""
    location:    str = ""
    dob:         str = ""
    nationality: str = ""
    summary:        str  = ""
    work:           list = field(default_factory=list)
    projects:       list = field(default_factory=list)
    education:      list = field(default_factory=list)
    skills:         dict = field(default_factory=dict)
    patents:        list = field(default_factory=list)
    awards:         list = field(default_factory=list)
    languages:      list = field(default_factory=list)
    certifications: list = field(default_factory=list)
    country_summaries: dict = field(default_factory=dict)
    country_work:      dict = field(default_factory=dict)
# =============================================================================
#  DEFAULT CANDIDATE DATA
# =============================================================================
def _default_data() -> ResumeData:
    return ResumeData(
        name        = "BALJINDER SINGH",
        email       = "baljindersinghcse@gmail.com",
        phone       = "+91 98158 11510",
        linkedin    = "linkedin.com/in/baljinder-singh-013b4311b",
        location    = "Bangalore, India",
        dob         = "20 March 1993",
        nationality = "Indian",
        summary = (
            "Results-driven Lead Full Stack Software Engineer with 9+ years building and scaling "
            "enterprise SaaS platforms. Expert in Java 17/Spring Boot backends, Angular/React frontends, "
            "and cloud-native development (GCP, AWS, Kubernetes). Delivered 35% performance gains, "
            "eliminated 40% of critical vulnerabilities, and built AI-powered automation tools. "
            "Patent holder for an AI-driven autonomous UI testing framework."
        ),
        work=[
            WorkEntry(
                company="OpenText", location="Bangalore, India",
                dates="Mar 2023 \u2013 Present",
                title="Lead Software Engineer",
                former_title="Senior Software Engineer",
                industry="Enterprise Information Management", company_size="15,000+ employees",
                bullets=[
                    "Architected CC4E Chatbot (Electron + Spring Boot MCP + FastAPI); 100+ daily conversations; MCP protocol & Playwright automation.",
                    "Led JATO V2\u2192V3 migration: AngularJS to Angular 17/Spring Boot; 50% capacity increase; Cloud Foundry\u2192GCP migration.",
                    "100% PSMQ compliance; 40% vulnerability reduction; 35% performance improvement via Redis caching; mentored 6+ engineers.",
                ],
            ),
            WorkEntry(
                company="Oracle", location="Bangalore, India",
                dates="Jan 2022 \u2013 Feb 2023",
                title="Senior Member of Technical Staff",
                industry="Enterprise Cloud Software", company_size="140,000+ employees",
                bullets=[
                    "SAR & Journals modules for FCCS cloud platform; audit-compliant RESTful services for Fortune 500 clients; 85%+ code coverage.",
                ],
            ),
            WorkEntry(
                company="OpenText", location="Bangalore, India",
                dates="Jan 2020 \u2013 Jan 2022",
                title="Senior Software Engineer",
                industry="Enterprise Information Management", company_size="15,000+ employees",
                bullets=[
                    "V2 platform on Cloud Foundry (12 microservices, 99.9% uptime); Saga pattern; 200+ enterprise customers.",
                ],
            ),
            WorkEntry(
                company="LTIMindtree", location="Bangalore, India",
                dates="Jul 2016 \u2013 Jan 2020",
                title="Senior Software Engineer", former_title="Software Engineer",
                industry="IT Services & Consulting", company_size="30,000+ employees",
                bullets=[
                    "Enterprise backend services (Java, Spring Boot); automated SLA dashboards and server health-checks; <4-hour production issue resolution.",
                ],
            ),
        ],
        projects=[
            WorkEntry(
                company="AI-Powered System Design Interview Platform",
                location="", dates="Full Stack (React 18 + Spring Boot 3.3.5)",
                title="Personal Project",
                bullets=[
                    "Full-stack: React 18 + Spring Boot 3.3.5 (Hexagonal Architecture) with LangChain4j, Ollama, Gemini API; WebSocket real-time sync & FSM interview flow.",
                    "Tech stack: PostgreSQL + Flyway, Redis caching, Prometheus/Grafana observability, TailwindCSS, React Flow.",
                ],
            ),
            WorkEntry(
                company="Basalt: AI Chat + RAG Platform",
                location="", dates="Full Stack (Java 21 + Spring Boot 3.2.5)",
                title="Personal Project",
                bullets=[
                    "Java 21 + Spring Boot 3.2.5 backend with Spring AI, streaming SSE; PDF RAG pipeline via pgvector + semantic chunking; Ollama & Pollinations.ai integration.",
                    "Angular 17 frontend with ngx-markdown; Docker Compose deployment; full-stack AI chat with local LLM inference.",
                ],
            ),
        ],
        education=[
            EduEntry(
                institution="Chandigarh University",
                degree="Bachelor of Engineering, Computer Science",
                dates="Jun 2012 \u2013 Jul 2016",
                location="India",
            ),
        ],
        skills={
            "Languages & Frameworks": "Java 17, Python, Spring Boot, FastAPI, Node.js, Hibernate, Angular, React, TypeScript, JavaScript, GraphQL, HTML5/CSS3",
            "Cloud & DevOps":         "AWS, Terraform, GCP, Kubernetes, Docker, Cloud Foundry, Okta, GitLab CI/CD, Jenkins, Maven, Git",
            "Data & Messaging":       "PostgreSQL, Oracle SQL, MySQL, Cassandra, XDB, DynamoDB, Redis, Kafka, RabbitMQ",
            "Architecture":           "Microservices, RESTful APIs, Saga Pattern, Event-Driven Design, Distributed Systems",
            "Testing & Quality":      "JUnit, Mockito, Karma, Integration Testing, E2E Testing, SonarQube",
            "AI & Automation":        "MCP Server/Client, Playwright, Ollama, Aviator (OpenText AI), Langchain, Langraph, LLM Orchestration",
            "Security & Compliance":  "Black Duck, Burp Suite, Fortify, PSMQ, OWASP Best Practices",
            "Design & Methodology":   "Figma (Design Reference), Agile/Scrum, Code Reviews, ADRs, Mentoring",
        },
        patents=[
            PatentEntry(
                title="Agentic AI-driven UI Testing Framework",
                year="2025",
                description="LLM-powered autonomous UI testing framework (MCP + Playwright), reducing manual testing by 70% and accelerating releases by 25%.",
            ),
        ],
        awards=[
            "Put Customers First (Feb 2026) \u2013 Resolved critical customer issues; led xPlore storage upgrade validation.",
            "Tackle Challenges Head On (Dec 2025) \u2013 Outstanding teamwork and collaborative contribution.",
            "Be Deserving of Trust (Oct 2023) \u2013 Resolved QatarGas escalation under pressure.",
            "Own the Outcome (Mar 2023) \u2013 3-year Supplier Exchange V1\u2192V2 migration; datacenter retirement & cost savings.",
            "Additional awards: Dec 2021, Mar 2021, Aug 2020 \u2013 Production support, escalation resolution, code coverage.",
        ],
        languages=[
            ("English",  "Professional (C1)"),
            ("Hindi",    "Native"),
            ("Punjabi",  "Native"),
        ],
    )
# =============================================================================
#  RESUME PARSER
# =============================================================================
class ResumeParser:
    @staticmethod
    def extract_text(file_path: str) -> str:
        ext = file_path.rsplit(".", 1)[-1].lower()
        if ext == "docx":
            return ResumeParser._from_docx(file_path)
        if ext == "pdf":
            return ResumeParser._from_pdf(file_path)
        if ext in ("txt", "md", "text"):
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        raise ValueError(f"Unsupported file type: .{ext}  (use .docx / .pdf / .txt)")
    @staticmethod
    def _from_docx(path: str) -> str:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    @staticmethod
    def _from_pdf(path: str) -> str:
        try:
            import fitz
        except ImportError:
            raise ImportError("PDF parsing requires PyMuPDF: pip install PyMuPDF")
        doc = fitz.open(path)
        return "\n".join(page.get_text() for page in doc)
# =============================================================================
#  GEMINI AI ENHANCER
# =============================================================================
class GeminiEnhancer:
    def __init__(self, api_key: str):
        if not api_key:
            raise ValueError("GEMINI_API_KEY is not set.")
        try:
            import google.generativeai as genai
        except ImportError:
            raise ImportError("Missing SDK: pip install google-generativeai")
        genai.configure(api_key=api_key)
        self._model = genai.GenerativeModel(GEMINI_MODEL)
    def _call(self, prompt: str) -> str:
        return self._model.generate_content(prompt).text.strip()
    def _parse_json(self, raw: str):
        text = raw.strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[1].rsplit("```", 1)[0]
        return json.loads(text.strip())
    def extract_structured_data(self, raw_text: str) -> ResumeData:
        schema = {
            "name": "", "email": "", "phone": "", "linkedin": "",
            "github": "", "website": "", "location": "", "dob": "", "nationality": "",
            "summary": "",
            "work": [{"company":"","location":"","dates":"","title":"","former_title":None,
                      "bullets":[],"industry":"","company_size":""}],
            "projects": [{"company":"","location":"","dates":"","title":"","bullets":[]}],
            "education": [{"institution":"","degree":"","dates":"","location":""}],
            "skills": {"Category": "val1, val2"},
            "patents": [{"title":"","year":"","description":""}],
            "awards": [],
            "languages": [{"language":"","proficiency":""}],
            "certifications": [],
        }
        prompt = (
            "Parse the following resume into the exact JSON schema. "
            "Return ONLY valid JSON - no markdown, no commentary.\n\n"
            f"Schema:\n{json.dumps(schema, indent=2)}\n\nResume:\n{raw_text}"
        )
        d = self._parse_json(self._call(prompt))
        return self._dict_to_resume_data(d)
    def _dict_to_resume_data(self, d: dict) -> ResumeData:
        def _work(items):
            return [WorkEntry(
                company=w.get("company",""), location=w.get("location",""),
                dates=w.get("dates",""), title=w.get("title",""),
                former_title=w.get("former_title"),
                bullets=w.get("bullets",[]),
                industry=w.get("industry"), company_size=w.get("company_size"),
            ) for w in (items or [])]
        def _edu(items):
            return [EduEntry(
                institution=e.get("institution",""), degree=e.get("degree",""),
                dates=e.get("dates",""), location=e.get("location"),
            ) for e in (items or [])]
        def _patents(items):
            return [PatentEntry(title=p.get("title",""), year=p.get("year",""),
                                description=p.get("description","")) for p in (items or [])]
        langs = [(l.get("language",""), l.get("proficiency",""))
                 for l in (d.get("languages") or [])]
        return ResumeData(
            name=d.get("name",""), email=d.get("email",""), phone=d.get("phone",""),
            linkedin=d.get("linkedin",""), github=d.get("github",""), website=d.get("website",""),
            location=d.get("location",""), dob=d.get("dob",""), nationality=d.get("nationality",""),
            summary=d.get("summary",""),
            work=_work(d.get("work",[])), projects=_work(d.get("projects",[])),
            education=_edu(d.get("education",[])), skills=d.get("skills",{}),
            patents=_patents(d.get("patents",[])), awards=d.get("awards",[]),
            languages=langs, certifications=d.get("certifications",[]),
        )
    def generate_country_summaries(self, data: ResumeData, countries: list, jd: str = "") -> dict:
        jd_ctx = f"\nTarget job description:\n{jd[:800]}" if jd else ""
        norms = {
            "canada":    "ATS-friendly, quantified wins, open to relocation to Canada, PR / work-permit mention",
            "usa":       "Power verbs, strong value proposition, metrics-first, concise",
            "uk":        "Formal British English, balance hard and soft skills, no age/nationality bias",
            "europe":    "Europass style, multilingual awareness, EU work authorisation mention",
            "japan":     "Formal tone, kaizen mindset, team harmony (wa), long-term commitment",
            "australia": "Friendly professional, cultural fit, right-to-work statement",
            "uae":       "International multicultural experience, achievements-first, UAE work-visa ready",
            "singapore": "Tech-forward, strong credentials, quantified outcomes",
        }
        results = {}
        for country in countries:
            prompt = (
                f"Write a 3-4 sentence professional summary for the following candidate, "
                f"tailored for the {country.upper()} job market.\n"
                f"Hiring norms: {norms.get(country, 'professional ATS-optimised')}\n"
                f"Name: {data.name}, Location: {data.location}\n"
                f"Current summary: {data.summary}\n"
                f"Top skills: {list(data.skills.values())[:3]}"
                f"{jd_ctx}\n\nReturn ONLY the summary text."
            )
            try:
                results[country] = self._call(prompt)
                print(f"  [AI] {country.upper()} summary OK")
            except Exception as exc:
                print(f"  [AI] {country.upper()} failed: {exc}")
        return results
    def enhance_bullets_for_country(self, bullets: list, country: str, jd: str = "") -> list:
        if not bullets:
            return bullets
        jd_ctx = f"\nJD keywords to weave in:\n{jd[:600]}" if jd else ""
        prompt = (
            f"Improve these resume bullets for the {country.upper()} market. "
            "Use stronger action verbs and quantify where possible."
            f"{jd_ctx}\n\nOriginal:\n" +
            "\n".join(f"- {b}" for b in bullets) +
            '\n\nReturn ONLY a JSON array of strings, same count as input.'
        )
        try:
            enhanced = self._parse_json(self._call(prompt))
            if isinstance(enhanced, list) and len(enhanced) == len(bullets):
                return enhanced
        except Exception:
            pass
        return bullets
    def enhance_all_work_for_country(self, work: list, country: str, jd: str = "") -> list:
        result = []
        for entry in work:
            new_e = deepcopy(entry)
            new_e.bullets = self.enhance_bullets_for_country(entry.bullets, country, jd)
            result.append(new_e)
        return result
# =============================================================================
#  DOCX HELPERS
# =============================================================================
def _margins(doc, top=0.5, bot=0.4, left=0.55, right=0.55):
    for s in doc.sections:
        s.top_margin = Inches(top); s.bottom_margin = Inches(bot)
        s.left_margin = Inches(left); s.right_margin = Inches(right)
def _heading_line(doc, color_hex="1A3C6E", space_before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1");   bot.set(qn("w:color"), color_hex)
    pBdr.append(bot); pPr.append(pBdr)
def _section(doc, text, color=NAVY, size=11, space_before=8, line_color="1A3C6E"):
    _heading_line(doc, color_hex=line_color, space_before=space_before)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = "Calibri"
def _bullet(doc, text, size=9.5, indent=0.25, after=1):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = "Calibri"
def _text(doc, text, size=9.5, bold=False, italic=False, color=None, after=2, before=0, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = "Calibri"
    if color: r.font.color.rgb = color
    return p
def _name_block(doc, name, subtitle=None):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run(name)
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY; r.font.name = "Calibri"
    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(4)
        rs = sub.add_run(subtitle)
        rs.font.size = Pt(11); rs.font.name = "Calibri"; rs.font.color.rgb = GRAY
def _contact_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line); r.font.size = Pt(9.5); r.font.name = "Calibri"
def _contact_line_link(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text); r.font.size = Pt(9.5); r.font.name = "Calibri"; r.font.color.rgb = LINK
def _contact_badge(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(9.5); r.font.name = "Calibri"; r.font.color.rgb = NAVY
def _company(doc, name, dates, title=None, former_title=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{name}  |  {dates}")
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.name = "Calibri"
    if title:
        pt = doc.add_paragraph()
        pt.paragraph_format.space_before = Pt(1); pt.paragraph_format.space_after = Pt(1)
        rt = pt.add_run(title)
        rt.italic = True; rt.font.size = Pt(10); rt.font.name = "Calibri"
        if former_title:
            rf = pt.add_run(f"  (formerly {former_title})")
            rf.italic = True; rf.font.size = Pt(10); rf.font.name = "Calibri"
            rf.font.color.rgb = GRAY
def _company_with_meta(doc, name, dates, title, meta=None, former_title=None):
    _company(doc, name, dates, title, former_title)
    if meta:
        pm = doc.add_paragraph()
        pm.paragraph_format.space_before = Pt(0); pm.paragraph_format.space_after = Pt(1)
        rm = pm.add_run(meta)
        rm.italic = True; rm.font.size = Pt(9); rm.font.name = "Calibri"; rm.font.color.rgb = GRAY
def _skill_line(doc, cat, val, size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(f"{cat}: "); r1.bold = True; r1.font.size = Pt(size); r1.font.name = "Calibri"
    r2 = p.add_run(val); r2.font.size = Pt(size); r2.font.name = "Calibri"
def _info_line(doc, label, value, link_color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(f"{label}: "); r1.bold = True; r1.font.size = Pt(9.5); r1.font.name = "Calibri"
    r2 = p.add_run(value); r2.font.size = Pt(9.5); r2.font.name = "Calibri"
    if link_color: r2.font.color.rgb = link_color
def _edu_row(doc, institution, degree, dates, tab_width=7.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(0)
    ts = p.paragraph_format.tab_stops
    ts.add_tab_stop(Inches(tab_width), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(institution); r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Calibri"
    r2 = p.add_run(f"\t{dates}"); r2.font.size = Pt(9.5); r2.font.name = "Calibri"; r2.font.color.rgb = GRAY
    _text(doc, degree, size=9.5, after=0)
def _patent_block(doc, patent: PatentEntry):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(f"Inventor \u2013 {patent.title}  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Calibri"
    r2 = p.add_run(f"(Filed & Completed, {patent.year})")
    r2.font.size = Pt(9.5); r2.font.name = "Calibri"
    _bullet(doc, patent.description)
def _work_section(doc, work_list, show_meta=False, get_bullets=None):
    for e in work_list:
        bullets = get_bullets(e) if get_bullets else e.bullets
        loc_part = f" | {e.location}" if e.location else ""
        if show_meta and (e.industry or e.company_size):
            meta = "  |  ".join(filter(None, [e.industry, e.company_size]))
            _company_with_meta(doc, f"{e.company}{loc_part}", e.dates, e.title, meta, e.former_title)
        else:
            _company(doc, f"{e.company}{loc_part}", e.dates, e.title, e.former_title)
        for b in bullets:
            _bullet(doc, b)
def _projects_section(doc, projects):
    for proj in projects:
        _company(doc, proj.company, proj.dates)
        for b in proj.bullets:
            _bullet(doc, b)
def _skills_section(doc, skills: dict):
    for cat, val in skills.items():
        _skill_line(doc, cat, val)
def _awards_section(doc, awards: list):
    for a in awards:
        _bullet(doc, a)
def _languages_section(doc, languages: list):
    for lang, level in languages:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        r1 = p.add_run(f"{lang}: "); r1.bold = True; r1.font.size = Pt(9.5); r1.font.name = "Calibri"
        r2 = p.add_run(level); r2.font.size = Pt(9.5); r2.font.name = "Calibri"
def _init_doc(top=0.5, bot=0.4, left=0.55, right=0.55):
    doc = Document()
    _margins(doc, top, bot, left, right)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    return doc
def _resolve_work(data: ResumeData, country: str):
    return data.country_work.get(country, data.work)
def _resolve_summary(data: ResumeData, country: str):
    return data.country_summaries.get(country, data.summary)
# =============================================================================
#  COUNTRY BUILDERS
# =============================================================================
def build_canada(data: ResumeData, out_dir: str):
    doc = _init_doc()
    _name_block(doc, data.name)
    _contact_block(doc, [f"{data.phone}  |  {data.email}  |  {data.location}"])
    if data.linkedin: _contact_line_link(doc, data.linkedin)
    _contact_badge(doc, "Open to Relocation to Canada  |  Willing to Obtain Canadian Work Permit / PR")
    _section(doc, "Professional Summary")
    _text(doc, _resolve_summary(data, "canada"), size=9.5, after=3)
    _section(doc, "Core Competencies")
    _skills_section(doc, data.skills)
    _section(doc, "Professional Experience")
    _work_section(doc, _resolve_work(data, "canada"))
    if data.projects:
        _section(doc, "Personal Projects"); _projects_section(doc, data.projects)
    if data.patents:
        _section(doc, "Innovation & Patent")
        for pat in data.patents: _patent_block(doc, pat)
    if data.awards:
        _section(doc, "Honours & Awards"); _awards_section(doc, data.awards)
    _section(doc, "Education")
    for e in data.education:
        loc = f", {e.location}" if e.location else ""
        _edu_row(doc, f"{e.institution}{loc}", e.degree, e.dates)
    _section(doc, "Languages")
    _text(doc, "  |  ".join(f"{l[0]} ({l[1]})" for l in data.languages), size=9.5, after=0)
    path = os.path.join(out_dir, f"{data.name.replace(' ','_')}_Canada.docx")
    doc.save(path); print(f"  [Canada]    {path}")
def build_usa(data: ResumeData, out_dir: str):
    doc = _init_doc()
    _name_block(doc, data.name)
    parts = [data.phone, data.email]
    if data.linkedin: parts.append(data.linkedin)
    if data.github:   parts.append(data.github)
    _contact_block(doc, ["  |  ".join(filter(None, parts)), data.location])
    _section(doc, "Professional Summary")
    _text(doc, _resolve_summary(data, "usa"), size=9.5, after=3)
    _section(doc, "Technical Skills")
    _skills_section(doc, data.skills)
    _section(doc, "Professional Experience")
    _work_section(doc, _resolve_work(data, "usa"))
    if data.projects:
        _section(doc, "Notable Projects"); _projects_section(doc, data.projects)
    if data.patents:
        _section(doc, "Patents & Innovation")
        for pat in data.patents: _patent_block(doc, pat)
    if data.certifications:
        _section(doc, "Certifications")
        for c in data.certifications: _bullet(doc, c)
    _section(doc, "Education")
    for e in data.education:
        loc = f", {e.location}" if e.location else ""
        _edu_row(doc, f"{e.institution}{loc}", e.degree, e.dates)
    if data.awards:
        _section(doc, "Awards & Recognition"); _awards_section(doc, data.awards[:3])
    path = os.path.join(out_dir, f"{data.name.replace(' ','_')}_USA.docx")
    doc.save(path); print(f"  [USA]       {path}")
def build_uk(data: ResumeData, out_dir: str):
    doc = _init_doc(left=0.6, right=0.6)
    _name_block(doc, data.name, subtitle="Curriculum Vitae")
    _contact_block(doc, [f"{data.phone}  |  {data.email}  |  {data.location}"])
    if data.linkedin: _contact_line_link(doc, data.linkedin)
    _section(doc, "Personal Profile")
    _text(doc, _resolve_summary(data, "uk"), size=9.5, after=3)
    _section(doc, "Key Skills")
    _skills_section(doc, data.skills)
    _section(doc, "Career History")
    _work_section(doc, _resolve_work(data, "uk"))
    if data.projects:
        _section(doc, "Personal Projects"); _projects_section(doc, data.projects)
    if data.patents:
        _section(doc, "Innovation & Patent")
        for pat in data.patents: _patent_block(doc, pat)
    _section(doc, "Education & Qualifications")
    for e in data.education:
        loc = f", {e.location}" if e.location else ""
        _edu_row(doc, f"{e.institution}{loc}", e.degree, e.dates, tab_width=6.4)
    _section(doc, "Languages")
    _languages_section(doc, data.languages)
    if data.awards:
        _section(doc, "Distinctions & Awards"); _awards_section(doc, data.awards)
    path = os.path.join(out_dir, f"{data.name.replace(' ','_')}_UK_CV.docx")
    doc.save(path); print(f"  [UK]        {path}")
def build_europe(data: ResumeData, out_dir: str):
    doc = _init_doc(left=0.6, right=0.6)
    _name_block(doc, data.name, subtitle="Curriculum Vitae")
    _section(doc, "Personal Information")
    _info_line(doc, "Email", data.email)
    _info_line(doc, "Phone", data.phone)
    _info_line(doc, "Location", f"{data.location}  \u2013  Open to relocation across Europe (EU Blue Card eligible)")
    if data.linkedin:    _info_line(doc, "LinkedIn",      data.linkedin, link_color=LINK)
    if data.dob:         _info_line(doc, "Date of Birth", data.dob)
    if data.nationality: _info_line(doc, "Nationality",   data.nationality)
    _section(doc, "Personal Profile")
    _text(doc, _resolve_summary(data, "europe"), size=9.5, after=3)
    _section(doc, "Key Skills & Competencies")
    _skills_section(doc, data.skills)
    _section(doc, "Work Experience")
    _work_section(doc, _resolve_work(data, "europe"))
    if data.projects:
        _section(doc, "Personal Projects"); _projects_section(doc, data.projects)
    if data.patents:
        _section(doc, "Innovation & Patent")
        for pat in data.patents: _patent_block(doc, pat)
    _section(doc, "Education")
    for e in data.education:
        loc = f", {e.location}" if e.location else ""
        _edu_row(doc, f"{e.institution}{loc}", e.degree, e.dates, tab_width=6.4)
    _section(doc, "Languages")
    _languages_section(doc, data.languages)
    if data.awards:
        _section(doc, "Honours & Awards"); _awards_section(doc, data.awards)
    _text(doc, "", size=6, after=0)
    _text(doc, "I hereby authorise the processing of my personal data pursuant to the European General Data Protection Regulation (EU) 2016/679 and applicable local legislation.", size=8, italic=True, color=GRAY, before=4, after=0)
    path = os.path.join(out_dir, f"{data.name.replace(' ','_')}_Europe_CV.docx")
    doc.save(path); print(f"  [Europe]    {path}")
def build_japan(data: ResumeData, out_dir: str):
    doc = _init_doc(left=0.6, right=0.6)
    _name_block(doc, data.name, subtitle="\u8077\u52d9\u7d4c\u6b74\u66f8  (Shokumu Keirekisho \u2013 Career Summary)")
    _section(doc, "Personal Details")
    _info_line(doc, "Full Name",       data.name.title())
    if data.dob:         _info_line(doc, "Date of Birth",    data.dob)
    if data.nationality: _info_line(doc, "Nationality",      data.nationality)
    _info_line(doc, "Current Location", data.location)
    _info_line(doc, "Visa Status", "Willing to obtain Engineer / Specialist in Humanities / International Services visa")
    _info_line(doc, "Email", data.email)
    _info_line(doc, "Phone", data.phone)
    if data.linkedin: _info_line(doc, "LinkedIn", data.linkedin, link_color=LINK)
    _section(doc, "Career Objective")
    _text(doc, _resolve_summary(data, "japan"), size=9.5, after=3)
    _section(doc, "Technical Skills  (\u6280\u8853\u30b9\u30ad\u30eb)")
    _skills_section(doc, data.skills)
    _section(doc, "Work Experience  (\u8077\u52d9\u7d4c\u6b74)")
    _work_section(doc, _resolve_work(data, "japan"), show_meta=True)
    if data.projects:
        _section(doc, "Personal Projects  (\u500b\u4eba\u30d7\u30ed\u30b8\u30a7\u30af\u30c8)")
        _projects_section(doc, data.projects)
    if data.patents:
        _section(doc, "Innovation & Patent  (\u7279\u8a31)")
        for pat in data.patents: _patent_block(doc, pat)
    _section(doc, "Education  (\u5b66\u6b74)")
    for e in data.education:
        loc = f", {e.location}" if e.location else ""
        _edu_row(doc, f"{e.institution}{loc}", e.degree, e.dates, tab_width=6.2)
    _section(doc, "Language Proficiency  (\u8a9e\u5b66\u529b)")
    langs = list(data.languages) + [("Japanese", "Willing to learn")]
    _languages_section(doc, langs)
    if data.awards:
        _section(doc, "Honours & Awards  (\u8868\u5f70)"); _awards_section(doc, data.awards)
    _section(doc, "Self-PR  (\u81ea\u5df1PR)")
    _text(doc, "I am a dedicated engineer who values quality, continuous learning, and team collaboration. I have consistently taken ownership of challenging projects, mentored team members, and proactively improved system reliability and security. I respect the importance of team harmony and am committed to adapting to and contributing to Japanese work culture. I am eager to bring my technical expertise and leadership to your organisation.", size=9.5, after=0)
    path = os.path.join(out_dir, f"{data.name.replace(' ','_')}_Japan.docx")
    doc.save(path); print(f"  [Japan]     {path}")
def build_australia(data: ResumeData, out_dir: str):
    doc = _init_doc()
    _name_block(doc, data.name)
    _contact_block(doc, [f"{data.phone}  |  {data.email}  |  {data.location}"])
    if data.linkedin: _contact_line_link(doc, data.linkedin)
    _contact_badge(doc, "Open to Relocation to Australia  |  Willing to Obtain TSS / Skilled Worker Visa")
    _section(doc, "Professional Summary")
    _text(doc, _resolve_summary(data, "australia"), size=9.5, after=3)
    _section(doc, "Technical Skills")
    _skills_section(doc, data.skills)
    _section(doc, "Work Experience")
    _work_section(doc, _resolve_work(data, "australia"))
    if data.projects:
        _section(doc, "Personal Projects"); _projects_section(doc, data.projects)
    if data.patents:
        _section(doc, "Innovation & Patent")
        for pat in data.patents: _patent_block(doc, pat)
    _section(doc, "Education")
    for e in data.education:
        loc = f", {e.location}" if e.location else ""
        _edu_row(doc, f"{e.institution}{loc}", e.degree, e.dates)
    _section(doc, "Languages")
    _languages_section(doc, data.languages)
    if data.awards:
        _section(doc, "Awards & Recognition"); _awards_section(doc, data.awards)
    path = os.path.join(out_dir, f"{data.name.replace(' ','_')}_Australia.docx")
    doc.save(path); print(f"  [Australia] {path}")
def build_uae(data: ResumeData, out_dir: str):
    doc = _init_doc()
    _name_block(doc, data.name)
    _contact_block(doc, [f"{data.phone}  |  {data.email}  |  {data.location}"])
    if data.linkedin: _contact_line_link(doc, data.linkedin)
    _contact_badge(doc, "Open to Relocation to UAE  |  Eligible for Employment Visa Sponsorship")
    _section(doc, "Professional Summary")
    _text(doc, _resolve_summary(data, "uae"), size=9.5, after=3)
    _section(doc, "Core Competencies")
    _skills_section(doc, data.skills)
    _section(doc, "Professional Experience")
    _work_section(doc, _resolve_work(data, "uae"))
    if data.projects:
        _section(doc, "Key Projects"); _projects_section(doc, data.projects)
    if data.patents:
        _section(doc, "Innovation & Patent")
        for pat in data.patents: _patent_block(doc, pat)
    _section(doc, "Education")
    for e in data.education:
        loc = f", {e.location}" if e.location else ""
        _edu_row(doc, f"{e.institution}{loc}", e.degree, e.dates)
    if data.nationality: _info_line(doc, "Nationality", data.nationality)
    _section(doc, "Languages"); _languages_section(doc, data.languages)
    if data.awards:
        _section(doc, "Achievements & Awards"); _awards_section(doc, data.awards)
    path = os.path.join(out_dir, f"{data.name.replace(' ','_')}_UAE.docx")
    doc.save(path); print(f"  [UAE]       {path}")
def build_singapore(data: ResumeData, out_dir: str):
    doc = _init_doc()
    _name_block(doc, data.name)
    _contact_block(doc, [f"{data.phone}  |  {data.email}  |  {data.location}"])
    if data.linkedin: _contact_line_link(doc, data.linkedin)
    if data.github:   _contact_line_link(doc, data.github)
    _contact_badge(doc, "Open to Relocation to Singapore  |  Eligible for Employment Pass / S-Pass")
    _section(doc, "Professional Summary")
    _text(doc, _resolve_summary(data, "singapore"), size=9.5, after=3)
    _section(doc, "Technical Skills")
    _skills_section(doc, data.skills)
    _section(doc, "Professional Experience")
    _work_section(doc, _resolve_work(data, "singapore"))
    if data.projects:
        _section(doc, "Key Projects"); _projects_section(doc, data.projects)
    if data.patents:
        _section(doc, "Innovation & Patent")
        for pat in data.patents: _patent_block(doc, pat)
    if data.certifications:
        _section(doc, "Certifications")
        for c in data.certifications: _bullet(doc, c)
    _section(doc, "Education")
    for e in data.education:
        loc = f", {e.location}" if e.location else ""
        _edu_row(doc, f"{e.institution}{loc}", e.degree, e.dates)
    if data.nationality: _info_line(doc, "Nationality", data.nationality)
    _section(doc, "Languages"); _languages_section(doc, data.languages)
    if data.awards:
        _section(doc, "Awards & Recognition"); _awards_section(doc, data.awards)
    path = os.path.join(out_dir, f"{data.name.replace(' ','_')}_Singapore.docx")
    doc.save(path); print(f"  [Singapore] {path}")
# =============================================================================
#  BUILDER REGISTRY
# =============================================================================
BUILDERS = {
    "canada":    build_canada,
    "usa":       build_usa,
    "uk":        build_uk,
    "europe":    build_europe,
    "japan":     build_japan,
    "australia": build_australia,
    "uae":       build_uae,
    "singapore": build_singapore,
}
# =============================================================================
#  CLI ENTRY POINT
# =============================================================================
def main():
    parser = argparse.ArgumentParser(
        description="AI-Powered Multi-Country Resume Generator",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  python generate_country_resumes.py\n"
            "  python generate_country_resumes.py --source resume.docx\n"
            "  python generate_country_resumes.py --source resume.pdf --jd jd.txt --ai-summaries\n"
            "  python generate_country_resumes.py --countries canada usa uk --enhance-bullets\n"
        ),
    )
    parser.add_argument("--source",          metavar="FILE",
                        help="Existing resume to parse (DOCX / PDF / TXT)")
    parser.add_argument("--jd",              metavar="FILE",
                        help="Job description file for keyword tailoring (TXT / MD)")
    parser.add_argument("--countries",       metavar="CC", nargs="+",
                        choices=SUPPORTED_COUNTRIES,
                        help=f"Target countries. Choices: {', '.join(SUPPORTED_COUNTRIES)}")
    parser.add_argument("--ai-summaries",    action="store_true",
                        help="Generate country-specific summaries via Gemini AI")
    parser.add_argument("--enhance-bullets", action="store_true",
                        help="Rewrite bullet points with Gemini AI per target country")
    parser.add_argument("--no-ai",           action="store_true",
                        help="Skip all AI calls")
    parser.add_argument("--output-dir",      metavar="DIR", default=OUTPUT_DIR,
                        help=f"Output folder (default: {OUTPUT_DIR})")
    parser.add_argument("--list-countries",  action="store_true",
                        help="List supported countries and exit")
    args = parser.parse_args()
    if args.list_countries:
        print("Supported countries:", ", ".join(SUPPORTED_COUNTRIES))
        return
    os.makedirs(args.output_dir, exist_ok=True)
    countries = args.countries or SUPPORTED_COUNTRIES
    # Step 1: Load data
    jd_text = ""
    if args.jd:
        with open(args.jd, "r", encoding="utf-8") as f:
            jd_text = f.read()
        print(f"Job description loaded: {args.jd}")
    data = _default_data()
    if args.source:
        print(f"Parsing source resume: {args.source}")
        raw_text = ResumeParser.extract_text(args.source)
        if not args.no_ai and GEMINI_API_KEY:
            print("Extracting structured data with Gemini AI...")
            try:
                ai = GeminiEnhancer(GEMINI_API_KEY)
                data = ai.extract_structured_data(raw_text)
                print(f"  Extracted: {data.name} | {len(data.work)} roles | {len(data.skills)} skill categories")
            except Exception as exc:
                print(f"  [WARN] AI extraction failed ({exc}) -- using built-in data as fallback.")
        else:
            print("  (AI disabled or GEMINI_API_KEY not set -- built-in data used)")
    # Step 2: AI enhancements
    if not args.no_ai and GEMINI_API_KEY:
        ai = GeminiEnhancer(GEMINI_API_KEY)
        if args.ai_summaries:
            print(f"Generating AI summaries for: {', '.join(countries)}")
            data.country_summaries = ai.generate_country_summaries(data, countries, jd_text)
        if args.enhance_bullets:
            print(f"Enhancing bullet points for: {', '.join(countries)}")
            for country in countries:
                print(f"  Enhancing {country.upper()}...")
                data.country_work[country] = ai.enhance_all_work_for_country(data.work, country, jd_text)
    elif not args.no_ai and not GEMINI_API_KEY:
        if args.ai_summaries or args.enhance_bullets:
            print("[WARN] GEMINI_API_KEY not set -- AI features skipped.")
    # Step 3: Generate DOCX
    print(f"\nGenerating resumes -> {args.output_dir}")
    for country in countries:
        try:
            BUILDERS[country](data, args.output_dir)
        except Exception as exc:
            print(f"  [ERROR] {country}: {exc}")
    n = len(countries)
    print(f"\nDone! {n} resume{'s' if n != 1 else ''} generated in {args.output_dir}")
if __name__ == "__main__":
    main()
